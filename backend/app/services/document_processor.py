import os
import uuid
from pathlib import Path
from typing import Dict, List, Any, Optional
import PyPDF2
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document
import re
from datetime import datetime

class DocumentProcessor:
    """
    Handles document processing including text extraction, OCR, and metadata extraction
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.txt', '.docx', '.doc', '.png', '.jpg', '.jpeg', '.bmp', '.tiff'}
    
    def process_document(self, file_path: str, original_name: str) -> Optional[Dict[str, Any]]:
        """
        Process a document and extract text, metadata, and structure
        
        Args:
            file_path: Path to the document file
            original_name: Original filename
            
        Returns:
            Dictionary containing processed document data
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Generate unique document ID
            doc_id = str(uuid.uuid4())
            
            # Extract text based on file type
            if file_extension == '.pdf':
                doc_data = self._process_pdf(file_path, doc_id, original_name)
            elif file_extension == '.txt':
                doc_data = self._process_text(file_path, doc_id, original_name)
            elif file_extension in {'.docx', '.doc'}:
                doc_data = self._process_word(file_path, doc_id, original_name)
            elif file_extension in {'.png', '.jpg', '.jpeg', '.bmp', '.tiff'}:
                doc_data = self._process_image(file_path, doc_id, original_name)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Add common metadata
            doc_data.update({
                'doc_id': doc_id,
                'original_name': original_name,
                'file_type': file_extension,
                'processed_at': datetime.now().isoformat(),
                'file_size': os.path.getsize(file_path)
            })
            
            return doc_data
            
        except Exception as e:
            print(f"Error processing document {original_name}: {str(e)}")
            return None
    
    def _process_pdf(self, file_path: str, doc_id: str, original_name: str) -> Dict[str, Any]:
        """Process PDF documents"""
        doc = fitz.open(file_path)
        pages = []
        full_text = ""
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text
            text = page.get_text()
            
            # If text is empty or very short, try OCR
            if len(text.strip()) < 50:
                try:
                    # Convert page to image and apply OCR
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    
                    # Save temporary image
                    temp_img_path = f"temp_page_{page_num}.png"
                    with open(temp_img_path, "wb") as f:
                        f.write(img_data)
                    
                    # Apply OCR
                    ocr_text = pytesseract.image_to_string(temp_img_path)
                    text = ocr_text if len(ocr_text.strip()) > len(text.strip()) else text
                    
                    # Clean up
                    os.remove(temp_img_path)
                    
                except Exception as e:
                    print(f"OCR failed for page {page_num}: {str(e)}")
            
            # Clean and structure text
            cleaned_text = self._clean_text(text)
            
            # Split into paragraphs
            paragraphs = self._split_into_paragraphs(cleaned_text)
            
            page_data = {
                'page_number': page_num + 1,
                'text': cleaned_text,
                'paragraphs': paragraphs,
                'word_count': len(cleaned_text.split())
            }
            
            pages.append(page_data)
            full_text += cleaned_text + "\n\n"
        
        doc.close()
        
        return {
            'pages': pages,
            'full_text': full_text.strip(),
            'total_pages': len(pages),
            'total_words': len(full_text.split()),
            'document_type': 'pdf'
        }
    
    def _process_text(self, file_path: str, doc_id: str, original_name: str) -> Dict[str, Any]:
        """Process text documents"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        cleaned_text = self._clean_text(text)
        paragraphs = self._split_into_paragraphs(cleaned_text)
        
        return {
            'pages': [{
                'page_number': 1,
                'text': cleaned_text,
                'paragraphs': paragraphs,
                'word_count': len(cleaned_text.split())
            }],
            'full_text': cleaned_text,
            'total_pages': 1,
            'total_words': len(cleaned_text.split()),
            'document_type': 'text'
        }
    
    def _process_word(self, file_path: str, doc_id: str, original_name: str) -> Dict[str, Any]:
        """Process Word documents"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            cleaned_text = self._clean_text(full_text)
            paragraphs = self._split_into_paragraphs(cleaned_text)
            
            return {
                'pages': [{
                    'page_number': 1,
                    'text': cleaned_text,
                    'paragraphs': paragraphs,
                    'word_count': len(cleaned_text.split())
                }],
                'full_text': cleaned_text,
                'total_pages': 1,
                'total_words': len(cleaned_text.split()),
                'document_type': 'word'
            }
            
        except Exception as e:
            print(f"Error processing Word document: {str(e)}")
            raise e
    
    def _process_image(self, file_path: str, doc_id: str, original_name: str) -> Dict[str, Any]:
        """Process image documents using OCR"""
        try:
            # Apply OCR to extract text
            text = pytesseract.image_to_string(file_path)
            cleaned_text = self._clean_text(text)
            paragraphs = self._split_into_paragraphs(cleaned_text)
            
            return {
                'pages': [{
                    'page_number': 1,
                    'text': cleaned_text,
                    'paragraphs': paragraphs,
                    'word_count': len(cleaned_text.split())
                }],
                'full_text': cleaned_text,
                'total_pages': 1,
                'total_words': len(cleaned_text.split()),
                'document_type': 'image'
            }
            
        except Exception as e:
            print(f"Error processing image with OCR: {str(e)}")
            raise e
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>\~\`]', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    def _split_into_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """Split text into paragraphs with metadata"""
        paragraphs = []
        
        # Split by double newlines or long single newlines
        para_splits = re.split(r'\n\s*\n|\n{2,}', text)
        
        for i, para_text in enumerate(para_splits):
            para_text = para_text.strip()
            if para_text and len(para_text) > 10:  # Skip very short paragraphs
                # Split into sentences for better citation
                sentences = self._split_into_sentences(para_text)
                
                paragraph_data = {
                    'paragraph_number': i + 1,
                    'text': para_text,
                    'sentences': sentences,
                    'word_count': len(para_text.split()),
                    'sentence_count': len(sentences)
                }
                
                paragraphs.append(paragraph_data)
        
        return paragraphs
    
    def _split_into_sentences(self, text: str) -> List[Dict[str, Any]]:
        """Split paragraph into sentences"""
        # Simple sentence splitting (can be improved with NLTK)
        sentence_endings = re.compile(r'[.!?]+\s+')
        sentences = []
        
        parts = sentence_endings.split(text)
        
        for i, sentence in enumerate(parts):
            sentence = sentence.strip()
            if sentence and len(sentence) > 5:
                sentences.append({
                    'sentence_number': i + 1,
                    'text': sentence,
                    'word_count': len(sentence.split())
                })
        
        return sentences
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from document"""
        file_stats = os.stat(file_path)
        
        metadata = {
            'file_size': file_stats.st_size,
            'created_at': datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            'file_extension': Path(file_path).suffix.lower()
        }
        
        # Try to extract additional PDF metadata
        if Path(file_path).suffix.lower() == '.pdf':
            try:
                doc = fitz.open(file_path)
                pdf_metadata = doc.metadata
                
                metadata.update({
                    'title': pdf_metadata.get('title', ''),
                    'author': pdf_metadata.get('author', ''),
                    'subject': pdf_metadata.get('subject', ''),
                    'creator': pdf_metadata.get('creator', ''),
                    'producer': pdf_metadata.get('producer', ''),
                    'creation_date': pdf_metadata.get('creationDate', ''),
                    'modification_date': pdf_metadata.get('modDate', '')
                })
                
                doc.close()
            except Exception as e:
                print(f"Could not extract PDF metadata: {str(e)}")
        
        return metadata